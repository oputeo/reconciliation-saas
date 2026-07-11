#!/usr/bin/env python3
"""SmartDelta + ReconFlow dual-platform presentation manual (PDF + Word)."""

from fpdf import FPDF
from fpdf.enums import XPos, YPos
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
from datetime import datetime

BASE = os.path.dirname(os.path.abspath(__file__))
OUT_DIR = os.path.join(BASE, "..", "deliverables", "smartdelta-delta")
os.makedirs(OUT_DIR, exist_ok=True)
TODAY = datetime.now().strftime("%d %B %Y")
PDF_OUT = os.path.join(OUT_DIR, "SmartDelta_ReconFlow_Presentation_Manual_July2026.pdf")
DOCX_OUT = os.path.join(OUT_DIR, "SmartDelta_ReconFlow_Presentation_Manual_July2026.docx")

NAVY = (0, 51, 102)
BLUE = (0, 94, 184)
GREEN = (0, 120, 60)
TEAL = (0, 105, 92)
DARK = (40, 40, 40)
MUTED = (90, 90, 90)
WHITE = (255, 255, 255)
LIGHT = (245, 249, 252)


class ManualPDF(FPDF):
    def __init__(self):
        super().__init__(format="A4")
        self.set_auto_page_break(auto=True, margin=14)

    def footer(self):
        self.set_y(-13)
        self.set_font("Helvetica", size=7)
        self.set_text_color(*MUTED)
        self.cell(0, 8, f"SmartDelta + ReconFlow Presentation Manual | {TODAY} | Page {self.page_no()}", align="C")

    def banner(self, text, fill=BLUE):
        self.set_fill_color(*fill)
        self.set_text_color(*WHITE)
        self.set_font("Helvetica", "B", 10)
        self.set_x(12)
        self.cell(0, 8, f"  {text}", fill=True, new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        self.ln(2)

    def body(self, text, size=9):
        self.set_font("Helvetica", size=size)
        self.set_text_color(*DARK)
        self.set_x(14)
        self.multi_cell(0, 4.8, text)
        self.ln(0.5)

    def bullet(self, text):
        self.set_font("Helvetica", size=9)
        self.set_text_color(*DARK)
        self.set_x(16)
        self.cell(4, 5, "-")
        self.multi_cell(0, 5, text)
        self.ln(0.3)

    def step(self, n, text):
        self.set_font("Helvetica", "B", 9)
        self.set_text_color(*NAVY)
        self.set_x(14)
        self.multi_cell(0, 5, f"Step {n}. {text}")
        self.ln(0.5)


def build_pdf():
    pdf = ManualPDF()
    pdf.add_page()
    pdf.set_fill_color(*NAVY)
    pdf.rect(0, 0, 210, 40, style="F")
    pdf.set_xy(12, 12)
    pdf.set_font("Helvetica", "B", 17)
    pdf.set_text_color(*WHITE)
    pdf.cell(0, 9, "SmartDelta Waste + ReconFlow", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.set_x(12)
    pdf.set_font("Helvetica", size=11)
    pdf.cell(0, 7, "Dual-Platform Presentation Manual - Delta State Finance Demo", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.set_x(12)
    pdf.set_font("Helvetica", size=9)
    pdf.cell(0, 6, f"Prepared {TODAY} | OEO Solution | Bi-weekly IGR / Pay Direct / Platform reconciliation", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.ln(10)

    pdf.banner("1. TWO-PLATFORM STORY (30-SECOND ELEVATOR)", TEAL)
    pdf.body(
        "SmartDelta proves collections happened (QR + photo + weight + Paystack). "
        "ReconFlow proves money landed in the right accounts (IGR 25%, PSP 65%, Platform 10%). "
        "Together they close the audit loop Finance requires."
    )

    pdf.banner("2. BEFORE THE MEETING - CHECKLIST (T-24 HOURS)", GREEN)
    for item in [
        "Railway: SmartDelta API + Frontend green (warm URLs 10 min before)",
        "ReconFlow: Log in as auditor; select tenant SmartDelta Waste - Delta State",
        "Upload 30k sample pack: C:\\ReconFlow-TestData\\SmartDelta-Delta\\chunks_platform then chunks_settlement",
        "Run setup-smartdelta-delta-tenant.mjs if ingest key not yet created",
        "Set SmartDelta API env: RECONFLOW_SUPABASE_URL, RECONFLOW_TENANT_ID, RECONFLOW_INGEST_KEY",
        "Print this manual + Finance handout v3",
        "Test logins: company_admin@smartdelta.ng, driver@smartdelta.ng / demo123",
    ]:
        pdf.bullet(item)

    pdf.banner("3. PRESENTATION RUN-OF-SHOW (45 MINUTES)", BLUE)
    pdf.step(1, "SmartDelta live (15 min) - Finance dashboard, Paystack payment, transaction ledger, 65/25/10 split.")
    pdf.step(2, "Driver proof (5 min) - Scan QR BIN, photo, weight, collection + KPI points.")
    pdf.step(3, "ReconFlow upload (5 min) - Show 30,000 rows ingested; Master Ledger count.")
    pdf.step(4, "Run bi-weekly reconciliation (5 min) - Period mode: Bi-weekly, BW06-2026.")
    pdf.step(5, "Anomalies walkthrough (10 min) - IGR shortfall, unmatched Paystack, fee leakage, duplicates.")
    pdf.step(6, "Close (5 min) - biweekly_summary.csv totals; Tranche 1 ask; Q&A.")

    pdf.add_page()
    pdf.banner("4. SMARTDELTA PLATFORM - KEY URLS & API", NAVY)
    pdf.body("Frontend: https://smartdelta-delta-production.up.railway.app")
    pdf.body("API: https://smartdelta-waste-delta-api-production.up.railway.app")
    pdf.body("Recon export: GET /api/finance/recon-export?period=BW06-2026&format=json")
    pdf.body("Recon push: POST /api/finance/recon-push { period, auto_reconcile: true }")
    pdf.body("Bi-week periods: GET /api/finance/recon-periods")

    pdf.banner("5. RECONFLOW - UPLOAD ORDER", TEAL)
    pdf.step(1, "Uploads -> Platform ledger: upload_ready/chunks_platform/ (Internal, qr_payment or smartdelta_paystack)")
    pdf.step(2, "Uploads -> Settlement: upload_ready/chunks_settlement/ (Settlement, bank_transfer)")
    pdf.step(3, "Reconciliation -> Period mode Bi-weekly -> BW06-2026 -> Run")
    pdf.step(4, "Anomalies -> filter High -> show rule tags (rule:EXC_UNMATCHED etc.)")
    pdf.step(5, "Executive tab -> match rate, variance exposure")

    pdf.banner("6. API INTEGRATION (LIVE PUSH DEMO)", GREEN)
    pdf.body("curl -X POST {SMARTDELTA_API}/api/finance/recon-push -H Content-Type:application/json -d {\"period\":\"BW06-2026\",\"auto_reconcile\":true}")
    pdf.body("This pushes live Paystack payments from SmartDelta DB into ReconFlow and optionally runs reconciliation.")

    pdf.banner("7. TALKING POINTS FOR FINANCE", TEAL)
    for item in [
        "State annual platform cost: NGN 0 under PPP - OEO bears hosting",
        "IGR 25% is accrued per payment and reconciled bi-weekly against Pay Direct credits",
        "Every anomaly has rule code, severity, and suggested action - not manual Excel",
        "30,000 transaction demo shows scale Ministry can expect at full rollout",
        "Back Audit available for historical unremitted funds (10-year lookback)",
    ]:
        pdf.bullet(item)

    pdf.banner("8. TROUBLESHOOTING", BLUE)
    pdf.bullet("ReconFlow 0% match: settlement chunks not uploaded - upload chunks_settlement")
    pdf.bullet("SmartDelta push 503: RECONFLOW_* env vars missing on Railway API")
    pdf.bullet("Driver submit fails: hard refresh; verify bin first; check API health")
    pdf.bullet("Cold start delay: warm both URLs 10 minutes before demo")

    pdf.banner("9. FILES REFERENCE", NAVY)
    pdf.bullet("Sample data: C:\\ReconFlow-TestData\\SmartDelta-Delta\\")
    pdf.bullet("Integration guide: deliverables/smartdelta-delta/SMARTDELTA_RECON_INTEGRATION_GUIDE.md")
    pdf.bullet("Tenant setup: scripts/setup-smartdelta-delta-tenant.mjs")
    pdf.bullet("Finance handout: SmartDelta_Finance_Handout_PermSec_Commissioner_July2026_v3.pdf")

    pdf.ln(4)
    pdf.set_font("Helvetica", "B", 9)
    pdf.set_text_color(*NAVY)
    pdf.set_x(14)
    pdf.cell(0, 6, "Opute Eric Opute (ACA) | OEO Solution | support@smartdelta.ng", new_x=XPos.LMARGIN, new_y=YPos.NEXT)

    pdf.output(PDF_OUT)
    print("PDF:", PDF_OUT)


def set_cell_shading(cell, fill_hex):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill_hex)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_shading(hdr[i], "005EB8")
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.color.rgb = RGBColor(255, 255, 255)
    for row_data in rows:
        row = table.add_row().cells
        for i, val in enumerate(row_data):
            row[i].text = str(val)
    doc.add_paragraph()


def build_docx():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("SmartDelta Waste + ReconFlow\nPresentation Manual")
    r.bold = True
    r.font.size = Pt(16)
    r.font.color.rgb = RGBColor(0, 51, 102)

    doc.add_heading("1. Two-Platform Story", level=1)
    doc.add_paragraph(
        "SmartDelta proves collections (QR, photo, weight, Paystack). "
        "ReconFlow proves remittances (IGR 25%, PSP 65%, Platform 10%). "
        "Together they close the Finance audit loop."
    )

    doc.add_heading("2. Pre-Meeting Checklist", level=1)
    for item in [
        "Warm SmartDelta URLs 10 minutes before",
        "Upload 30k ReconFlow sample data (platform + settlement chunks)",
        "Configure RECONFLOW_* env on SmartDelta API",
        "Select SmartDelta tenant in ReconFlow",
        "Test demo logins",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("3. Run-of-Show (45 min)", level=1)
    add_table(doc, ["Step", "Time", "Action"], [
        ["1", "15 min", "SmartDelta Finance dashboard + Paystack payment + ledger"],
        ["2", "5 min", "Driver QR collection demo"],
        ["3", "5 min", "ReconFlow upload 30k rows"],
        ["4", "5 min", "Bi-weekly reconciliation BW06-2026"],
        ["5", "10 min", "Anomalies: IGR shortfall, fee leakage, unmatched"],
        ["6", "5 min", "Close + Q&A"],
    ])

    doc.add_heading("4. SmartDelta URLs", level=1)
    doc.add_paragraph("Frontend: https://smartdelta-delta-production.up.railway.app")
    doc.add_paragraph("API: https://smartdelta-waste-delta-api-production.up.railway.app")
    doc.add_paragraph("GET /api/finance/recon-export?period=BW06-2026")
    doc.add_paragraph("POST /api/finance/recon-push with auto_reconcile: true")

    doc.add_heading("5. ReconFlow Upload Order", level=1)
    for item in [
        "chunks_platform/ -> Internal (qr_payment)",
        "chunks_settlement/ -> Settlement (bank_transfer)",
        "Reconciliation -> Bi-weekly -> BW06-2026",
        "Review Anomalies and Executive dashboard",
    ]:
        doc.add_paragraph(item, style="List Number")

    doc.add_heading("6. Troubleshooting", level=1)
    add_table(doc, ["Issue", "Fix"], [
        ["0% match rate", "Upload settlement chunks"],
        ["Push 503", "Set RECONFLOW env vars on Railway"],
        ["Driver submit fail", "Hard refresh; verify bin; check API health"],
        ["Slow first load", "Warm URLs before meeting"],
    ])

    doc.add_paragraph("Opute Eric Opute (ACA) | OEO Solution")
    doc.save(DOCX_OUT)
    print("DOCX:", DOCX_OUT)


if __name__ == "__main__":
    build_pdf()
    build_docx()
    print("Done.")